"""
Word Generator Tool
Converts HTML content into a Microsoft Word (.docx) document.
Supports headings, paragraphs, and lists.
"""

import datetime
from pathlib import Path
from typing import List, Dict
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

class WordGenerator:
    def __init__(self, output_dir: str = "generated_docs/word"):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def _parse_html(self, html: str) -> BeautifulSoup:
        return BeautifulSoup(html, "html.parser")

    def generate_docx(self, html_content: str, filename: str = None) -> str:
        """
        Generates a DOCX file from the provided HTML content.
        
        Args:
            html_content: HTML string.
            filename: Output filename.
            
        Returns:
            Absolute path to the generated DOCX file.
        """
        if not filename:
            timestamp = datetime.datetime.now().strftime("%Y%m%d-%H%M%S")
            filename = f"document_{timestamp}.docx"
            
        if not filename.endswith(".docx"):
            filename += ".docx"

        output_path = self.output_dir / filename
        
        document = Document()
        soup = self._parse_html(html_content)
        
        # Iterate over all tags in order
        for element in soup.find_all(["h1", "h2", "h3", "h4", "p", "li", "ul", "ol"]):
            tag_name = element.name
            text = element.get_text(strip=True)
            
            if not text:
                continue

            if tag_name.startswith("h"):
                # Headings
                level = int(tag_name[1])
                # Mapping HTML h1-h4 to Word Heading 1-4
                # Note: 'Title' style is often used for h1, but let's stick to Headings for structure
                heading = document.add_heading(text, level=level)
                
            elif tag_name == "p":
                # Paragraphs
                p = document.add_paragraph(text)
                
            elif tag_name == "li":
                # List items
                # Check parent to decide bullet vs number
                parent = element.parent.name if element.parent else "ul"
                style = 'List Bullet' if parent == 'ul' else 'List Number'
                document.add_paragraph(text, style=style)
                
        document.save(str(output_path))
        return str(output_path.absolute())

if __name__ == "__main__":
    generator = WordGenerator()
    html = """
    <h1>Project Report</h1>
    <h2>1. Introduction</h2>
    <p>This is a generated report using N-T-AI Office Suite.</p>
    <h2>2. Features</h2>
    <ul>
        <li>Fast generation</li>
        <li>Clean formatting</li>
    </ul>
    """
    print(generator.generate_docx(html))
